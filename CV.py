from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn


def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi',
                              'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing',
                              'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
                              'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '808080')
    pBdr.append(bottom)


def create_professional_cv():
    # Create new document
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Header with name
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run('UNNATI TALREJA')
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.name = 'Arial'
    name_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add line after name
    add_horizontal_line(header)

    # Contact information
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = [
        '+91 9325134574',
        'unnatigot7@gmail.com',
        'Bengaluru, Karnataka'
    ]
    contact_text = ' | '.join(contact_info)
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(11)
    contact_run.font.name = 'Arial'

    # Add space
    doc.add_paragraph()

    # Professional Summary
    summary_heading = doc.add_paragraph()
    summary_heading_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
    summary_heading_run.font.size = Pt(14)
    summary_heading_run.font.bold = True
    summary_heading_run.font.name = 'Arial'
    summary_heading_run.font.color.rgb = RGBColor(0, 0, 0)
    add_horizontal_line(summary_heading)

    summary_text = doc.add_paragraph()
    summary_content = ("Ambitious BBA student with demonstrated experience in project management, strategic planning, "
                       "digital marketing, and data analysis. Proficient in Excel, market research, AI tools, and "
                       "business consulting simulations. Proven ability in team coordination, adaptability, and "
                       "cross-functional communication. Seeking internship opportunities to apply academic knowledge "
                       "and gain practical business experience.")
    summary_run = summary_text.add_run(summary_content)
    summary_run.font.size = Pt(11)
    summary_run.font.name = 'Arial'
    summary_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Education
    education_heading = doc.add_paragraph()
    education_heading_run = education_heading.add_run('EDUCATION')
    education_heading_run.font.size = Pt(14)
    education_heading_run.font.bold = True
    education_heading_run.font.name = 'Arial'
    add_horizontal_line(education_heading)

    # Current Education
    current_edu = doc.add_paragraph()
    current_edu_run = current_edu.add_run('Bachelor of Business Administration (BBA)')
    current_edu_run.font.size = Pt(12)
    current_edu_run.font.bold = True
    current_edu_run.font.name = 'Arial'

    current_details = doc.add_paragraph()
    current_details.paragraph_format.left_indent = Inches(0.25)
    details_run = current_details.add_run('Alliance University, Bengaluru | 2023 - Present (Final Year)')
    details_run.font.size = Pt(11)
    details_run.font.name = 'Arial'
    details_run.italic = True

    # Previous Education
    prev_edu = doc.add_paragraph()
    prev_edu_run = prev_edu.add_run('Post Graduate Certificate in Management (PGCM)')
    prev_edu_run.font.size = Pt(12)
    prev_edu_run.font.bold = True
    prev_edu_run.font.name = 'Arial'

    prev_details = doc.add_paragraph()
    prev_details.paragraph_format.left_indent = Inches(0.25)
    prev_details_run = prev_details.add_run('Vidya Niketan | 2020 - 2022')
    prev_details_run.font.size = Pt(11)
    prev_details_run.font.name = 'Arial'
    prev_details_run.italic = True

    # Key Skills
    skills_heading = doc.add_paragraph()
    skills_heading_run = skills_heading.add_run('CORE COMPETENCIES')
    skills_heading_run.font.size = Pt(14)
    skills_heading_run.font.bold = True
    skills_heading_run.font.name = 'Arial'
    add_horizontal_line(skills_heading)

    # Create skills in columns
    skills_categories = [
        ('Business & Strategy:', ['Strategic Business Planning', 'Project Management', 'Market Research & Analysis']),
        ('Technical Skills:', ['Advanced Excel (100%)', 'AI & Digital Literacy', 'Data Analysis']),
        ('Marketing & Communication:',
         ['Digital Marketing', 'Social Media Management', 'Cross-functional Communication']),
        ('Leadership & Soft Skills:', ['Team Coordination', 'Problem-solving', 'Adaptability', 'Interpersonal Skills'])
    ]

    for category, skills in skills_categories:
        category_para = doc.add_paragraph()
        category_para.paragraph_format.left_indent = Inches(0.25)
        category_run = category_para.add_run(category)
        category_run.font.size = Pt(11)
        category_run.font.bold = True
        category_run.font.name = 'Arial'

        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.left_indent = Inches(0.5)
        skills_text = ' • '.join(skills)
        skills_run = skills_para.add_run(skills_text)
        skills_run.font.size = Pt(10)
        skills_run.font.name = 'Arial'

    # Academic Projects & Certifications
    projects_heading = doc.add_paragraph()
    projects_heading_run = projects_heading.add_run('ACADEMIC PROJECTS & CERTIFICATIONS')
    projects_heading_run.font.size = Pt(14)
    projects_heading_run.font.bold = True
    projects_heading_run.font.name = 'Arial'
    add_horizontal_line(projects_heading)

    # Capstone Project
    capstone_title = doc.add_paragraph()
    capstone_title.paragraph_format.left_indent = Inches(0.25)
    capstone_title_run = capstone_title.add_run('Strategic Business Plan Development - Capstone Project')
    capstone_title_run.font.size = Pt(12)
    capstone_title_run.font.bold = True
    capstone_title_run.font.name = 'Arial'

    capstone_desc = doc.add_paragraph()
    capstone_desc.paragraph_format.left_indent = Inches(0.5)
    capstone_text = ("Led comprehensive capstone project creating strategic business plan for startup venture. "
                     "Applied integrated concepts from Finance, HR, Consumer Behavior, Operations, and International Business. "
                     "Conducted market research, financial modeling, competitive analysis, and developed scalable, "
                     "financially viable business model with global expansion strategy.")
    capstone_run = capstone_desc.add_run(capstone_text)
    capstone_run.font.size = Pt(11)
    capstone_run.font.name = 'Arial'
    capstone_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Professional Training & Simulations
    training_title = doc.add_paragraph()
    training_title.paragraph_format.left_indent = Inches(0.25)
    training_title_run = training_title.add_run('Professional Training & Industry Simulations')
    training_title_run.font.size = Pt(12)
    training_title_run.font.bold = True
    training_title_run.font.name = 'Arial'

    # Coursera Certifications
    coursera_para = doc.add_paragraph()
    coursera_para.paragraph_format.left_indent = Inches(0.5)
    coursera_text = (
        "• Coursera Certifications: Financial Accounting, Accounting Cycle, Journals & Ledgers (UC Irvine, 98-100%), "
        "Management and Microeconomics (96-100%), Market Research (IE, 92%), Quality Management (94%), Excel (100%)")
    coursera_run = coursera_para.add_run(coursera_text)
    coursera_run.font.size = Pt(10)
    coursera_run.font.name = 'Arial'

    # Google & Forage Simulations
    simulations_para = doc.add_paragraph()
    simulations_para.paragraph_format.left_indent = Inches(0.5)
    simulations_text = ("• Google & Forage Professional Simulations: Consulting (BCG), Project Management (CBRE), "
                        "Finance (Citi), Digital Marketing, Analytics, and Project Management certifications. "
                        "Gained practical experience in strategy development, operations, and data-driven decision making.")
    simulations_run = simulations_para.add_run(simulations_text)
    simulations_run.font.size = Pt(10)
    simulations_run.font.name = 'Arial'

    # Additional Information
    additional_heading = doc.add_paragraph()
    additional_heading_run = additional_heading.add_run('ADDITIONAL INFORMATION')
    additional_heading_run.font.size = Pt(14)
    additional_heading_run.font.bold = True
    additional_heading_run.font.name = 'Arial'
    add_horizontal_line(additional_heading)

    additional_items = [
        "Languages: English (Fluent), Hindi (Native)",
        "Availability: Immediate for internship opportunities",
        "Interests: Business Strategy, Market Research, Digital Innovation, Data Analytics"
    ]

    for item in additional_items:
        item_para = doc.add_paragraph()
        item_para.paragraph_format.left_indent = Inches(0.25)
        item_run = item_para.add_run(f"• {item}")
        item_run.font.size = Pt(11)
        item_run.font.name = 'Arial'

    # Save the document
    filename = "Unnati_Talreja_Professional_CV.docx"
    doc.save(filename)
    print(f"Professional CV generated successfully: {filename}")
    return filename


# Run the function to create CV
if __name__ == "__main__":
    create_professional_cv()